from docx import Document

document = Document()

document.add_heading('Test Title', 0)

paragraph = document.add_paragraph('paragraph test')
paragraph.add_run('bold').bold = True
paragraph.add_run(' and ')
paragraph.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)

document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('unordered list', style='List Bullet')
document.add_paragraph('ordered list', style='List Number')

# document.add_picture('test.png', width=Inches(1.25))

records = (
    (1, '100', 'java'),
    (2, '101', 'c'),
    (3, '102', 'python')
)

table = document.add_table(rows=1, cols=3)
cells = table.rows[0].cells
cells[0].text = 'first'
cells[1].text = 'second'
cells[2].text = 'third'

for i, j, k in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = j
    row_cells[2].text = k

document.add_page_break()

document.save('test.docx')

f = open('test.docx', 'rb')
document = Document(f)
f.close()